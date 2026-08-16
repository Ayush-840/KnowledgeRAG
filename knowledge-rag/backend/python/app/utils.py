import csv
import os
import re
from collections import Counter
from typing import Callable, List, Optional, Tuple

from pypdf import PdfReader
import pdfplumber

from rank_bm25 import BM25Okapi

from .dependencies import get_session_vectors, persist_bm25
from .embedders import get_embedder

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".md", ".csv"}
MAX_CSV_ROWS = int(os.getenv("MAX_CSV_ROWS", "50000"))

# ----- Text extraction -----

# A line that is just a page number ("3", "Page 4", "— 7 —", "· 3 ·", with
# optional surrounding ornaments/spacing).
_PAGE_NUMBER_LINE = re.compile(r"^\s*(?:page\s+)?[.·—–-]*\s*\d{1,4}\s*[.·—–-]*\s*$", re.IGNORECASE)
# Glue pattern: a short numeric prefix fused directly onto the next word with no
# whitespace at the very start of a page (the "04Build" artifact). Conservative:
# only 1-3 digits followed by a capital letter — plain numbers in body sentences
# ("40GB", "2024 revenue") are mid-line or lowercase and are left untouched.
_GLUE_PREFIX = re.compile(r"^(\d{1,3})(?=[A-Z][a-z])")


def _normalize_repeat_key(line: str) -> str:
    """Collapse a line to a stable key ignoring digits/punctuation, so running
    footers like 'Aurora Labs — 4' and 'Aurora Labs — 5' count as the same line."""
    return re.sub(r"[\d\W_]+\s*", " ", line).strip().lower()


def detect_repeated_lines(pages: List[Tuple[int, str]]) -> set:
    """Return raw lines that recur near-identically in the top/bottom of >50% of
    pages — almost certainly running headers/footers (boilerplate), not content.
    Comparison uses a digit/punctuation-stripped key so footers containing page
    numbers still match across pages.
    """
    n = len(pages)
    if n < 3:
        return set()
    threshold = max(2, int(n * 0.5))
    top_keys, bottom_keys = Counter(), Counter()
    top_examples, bottom_examples = {}, {}
    for _, text in pages:
        lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
        if not lines:
            continue
        for ln in lines[:2]:
            key = _normalize_repeat_key(ln)
            if key:
                top_keys[key] += 1
                top_examples.setdefault(key, ln)
        for ln in lines[-2:]:
            key = _normalize_repeat_key(ln)
            if key:
                bottom_keys[key] += 1
                bottom_examples.setdefault(key, ln)
    repeated = set()
    for counter, examples in ((top_keys, top_examples), (bottom_keys, bottom_examples)):
        for key, count in counter.items():
            if count >= threshold and len(key) > 3:
                repeated.add(examples[key])
    return repeated


def clean_page_text(page_text: str, repeated_lines: set = None) -> str:
    """Strip PDF layout artifacts from one page of extracted text:

    - standalone page-number lines ("3", "Page 4")
    - running headers/footers that recur across pages (matched by normalized key
      so footers with varying page numbers are still caught)
    - the glue pattern: a short numeric prefix fused onto the first word of the
      page ("04Build") — the digits are dropped, the word is kept.

    Only the top/bottom lines of the page are eligible for header/footer removal,
    and glue-stripping is limited to the very first content line, so real numbers
    inside body text are never touched.
    """
    repeated_lines = repeated_lines or set()
    repeated_keys = {_normalize_repeat_key(ln) for ln in repeated_lines}
    lines = page_text.split("\n")
    cleaned = []
    glue_checked = False  # apply the glue-strip to the first surviving content line
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            cleaned.append(line)
            continue
        # Standalone page-number line, anywhere on the page
        if _PAGE_NUMBER_LINE.match(stripped):
            continue
        # Running header/footer: only the top 2 / bottom 2 lines of the page
        is_edge = i <= 1 or i >= len(lines) - 2
        if is_edge and _normalize_repeat_key(stripped) in repeated_keys:
            continue
        # Glue pattern: strip a short numeric prefix fused onto the first
        # surviving content line ("04Build" -> "Build"). The header/footer and
        # page-number checks above run first, so a header line that precedes the
        # content doesn't consume the glue-strip.
        if not glue_checked:
            glue_checked = True
            m = _GLUE_PREFIX.match(line)
            if m:
                line = line[m.end():]
                stripped = line.strip()
        cleaned.append(line)
    return "\n".join(cleaned).strip()


def clean_extracted_pages(pages: List[Tuple[int, str]]) -> List[Tuple[int, str]]:
    """Post-process every page of an extracted PDF: detect repeated
    headers/footers across pages first, then clean each page."""
    repeated = detect_repeated_lines(pages)
    return [(pno, clean_page_text(text, repeated)) for pno, text in pages]


def extract_text_from_pdf(file_path: str) -> List[Tuple[int, str]]:
    """Extract text from each page of a PDF, then strip layout artifacts
    (page numbers, running headers/footers, glued numeric prefixes).
    Returns a list of (page_number, page_text).
    """
    pages = []
    # Try pypdf first
    try:
        reader = PdfReader(file_path)
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append((i + 1, text))
    except Exception:
        # Fallback to pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                pages.append((i + 1, text))
    return clean_extracted_pages(pages)


def extract_text_from_txt(file_path: str) -> List[Tuple[int, str]]:
    """Treat a plain text file as a single page.
    Returns a list with one tuple (1, content).
    """
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()
    return [(1, content)]


def extract_blocks_from_docx(file_path: str) -> List[dict]:
    """Extract paragraphs from a .docx, preserving heading structure.
    Returns a list of {"text": str, "heading": Optional[str]} — a heading paragraph
    becomes a block with heading set and empty text.
    """
    from docx import Document

    doc = Document(file_path)
    blocks = []
    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = getattr(para.style, "name", None)
        style = (style_name or "").lower()
        if style.startswith("heading"):
            blocks.append({"text": "", "heading": text})
        else:
            blocks.append({"text": text, "heading": None})
    return blocks


def extract_blocks_from_markdown(file_path: str) -> List[dict]:
    """Extract a .md into paragraph blocks; heading lines (#..######) become
    heading blocks so the structure-aware chunker can split along headings.
    """
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()
    blocks = []
    for para in re.split(r"\n\s*\n", content):
        para = para.strip()
        if not para:
            continue
        lines = para.splitlines()
        m = re.match(r"^(#{1,6})\s+(.*)", lines[0])
        if m:
            blocks.append({"text": "", "heading": m.group(2).strip()})
            body = "\n".join(lines[1:]).strip()
            if body:
                blocks.append({"text": body, "heading": None})
        else:
            blocks.append({"text": para, "heading": None})
    return blocks


def extract_csv_chunks(file_path: str, max_words: int = 500) -> Tuple[List[Tuple[str, int, int]], int]:
    """Chunk a .csv row-wise: each chunk keeps the header row and groups data rows
    until max_words is reached, so chunks stay tabular and coherent.
    Returns ([(chunk_text, first_data_row, last_data_row)], data_row_count) — row
    numbers are 1-indexed file rows (header is row 1).
    """
    with open(file_path, newline="", encoding="utf-8", errors="ignore") as f:
        rows = [row for row in csv.reader(f) if any(cell.strip() for cell in row)]
    if not rows:
        raise ValueError("CSV file is empty")
    header = rows[0]
    body = rows[1:]
    if len(body) > MAX_CSV_ROWS:
        raise ValueError(f"CSV exceeds row limit of {MAX_CSV_ROWS} data rows")

    header_text = ", ".join(header)
    chunks = []
    current = []
    current_words = 0
    current_start = 2  # data rows start at file row 2 (header is row 1)
    for file_row, row in enumerate(body, start=2):
        row_text = ", ".join(row)
        words = len(row_text.split())
        if current and current_words + words > max_words:
            chunks.append(("\n".join([header_text, *current]), current_start, file_row - 1))
            current, current_words = [], 0
            current_start = file_row
        current.append(row_text)
        current_words += words
    if current:
        chunks.append(("\n".join([header_text, *current]), current_start, len(body) + 1))

    return chunks, len(body)

# ----- Chunking -----

def chunk_text(page_text: str, max_words: int = 500, overlap: int = 100) -> List[str]:
    """Fixed-size chunking: split page text into overlapping word windows.
    Returns a list of chunk strings.
    """
    words = page_text.split()
    chunks = []
    i = 0
    step = max(max_words - overlap, 1)  # never let overlap >= max_words stall the loop
    while i < len(words):
        chunk = words[i : i + max_words]
        chunks.append(" ".join(chunk))
        i += step
    return chunks


def chunk_text_structure_aware(page_text: str, max_words: int = 500, overlap: int = 100) -> List[str]:
    """Structure-aware chunking: respect paragraph boundaries (blank-line separated).
    Paragraphs are grouped into chunks up to max_words; a single oversized paragraph
    is split with the fixed-size splitter. Falls back to fixed-size chunking when the
    text has no paragraph breaks (unstructured text).
    """
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", page_text) if p.strip()]
    if not paragraphs:
        return chunk_text(page_text, max_words, overlap)

    chunks = []
    current: List[str] = []
    current_len = 0
    for para in paragraphs:
        para_words = para.split()
        if current_len + len(para_words) <= max_words:
            current.append(para)
            current_len += len(para_words)
        else:
            if current:
                chunks.append("\n\n".join(current))
            if len(para_words) > max_words:
                # Oversized paragraph: split fixed-size, then continue fresh
                chunks.extend(chunk_text(para, max_words, overlap))
                current, current_len = [], 0
            else:
                current, current_len = [para], len(para_words)
    if current:
        chunks.append("\n\n".join(current))
    return chunks


def chunk_blocks(blocks: List[dict], max_words: int = 500, overlap: int = 100, structure_aware: bool = False) -> List[str]:
    """Chunk block lists (docx/md). structure_aware: heading blocks start a new chunk
    (the heading is carried into the chunk); body blocks group up to max_words, with
    oversized blocks split fixed-size. fixed: flatten all blocks and word-split.
    """
    if not structure_aware:
        full = "\n\n".join(b["text"] for b in blocks if b["text"])
        return chunk_text(full, max_words, overlap)

    chunks = []
    current: List[str] = []
    current_len = 0
    for b in blocks:
        heading = b.get("heading")
        text = b["text"]
        if heading:
            if current:
                chunks.append("\n\n".join(current))
                current, current_len = [], 0
            current.append(heading)
            current_len = len(heading.split())
            continue
        if not text:
            continue
        words = text.split()
        if current_len + len(words) <= max_words:
            current.append(text)
            current_len += len(words)
        else:
            if current:
                chunks.append("\n\n".join(current))
                current, current_len = [], 0
            if len(words) > max_words:
                chunks.extend(chunk_text(text, max_words, overlap))
            else:
                current, current_len = [text], len(words)
    if current:
        chunks.append("\n\n".join(current))
    return chunks


CHUNKERS = {
    "fixed": chunk_text,
    "structure_aware": chunk_text_structure_aware,
}

# ----- Document titles -----

# Titles that are "empty/generic" and shouldn't be shown instead of a filename.
_GENERIC_TITLES = {"", "untitled", "untitled document", "document", "pdf", "microsoft word", "microsoft word document", "new document"}


def extract_document_title(filename: str, file_path: str, pages: Optional[List[Tuple[int, str]]] = None, blocks: Optional[List[dict]] = None) -> Optional[str]:
    """Best-effort document title for display purposes:

    - .pdf  — PDF /Title metadata
    - .docx — core-properties title
    - .md   — first H1 heading
    - .txt  — first line, when it's short enough to plausibly be a title
    - .csv  — none (chunks stay tabular)

    Returns None when no non-generic title exists, so callers fall back to the
    user's original upload filename.
    """
    ext = os.path.splitext(filename)[1].lower()
    title = None
    if ext == ".pdf":
        try:
            meta = PdfReader(file_path).metadata or {}
            title = meta.get("/Title")
        except Exception:  # noqa: BLE001 - title is best-effort
            title = None
    elif ext == ".docx":
        try:
            from docx import Document
            title = Document(file_path).core_properties.title
        except Exception:  # noqa: BLE001
            title = None
    elif ext == ".md" and blocks:
        for b in blocks:
            heading = (b.get("heading") or "").strip()
            if heading:
                title = heading
                break
    elif ext == ".txt" and pages:
        first = (pages[0][1] or "").strip().splitlines()
        if first:
            candidate = re.sub(r"^\s*#+\s*", "", first[0]).strip()
            if 0 < len(candidate) <= 80 and not _PAGE_NUMBER_LINE.match(candidate):
                title = candidate

    if not title:
        return None
    title = title.strip()
    if title.lower() in _GENERIC_TITLES or len(title) > 120:
        return None
    return title


# ----- Embedding -----

def embed_chunks(chunks: List[str]):
    """Return a list of dense vector embeddings for the given chunks."""
    return get_embedder().embed(chunks)

# ----- BM25 -----

def build_bm25_index(chunks: List[str]):
    """Create a BM25Okapi index from tokenized chunks.
    Returns the BM25 object and the tokenized corpus (list of list of tokens).
    """
    tokenized = [c.lower().split() for c in chunks]
    bm25 = BM25Okapi(tokenized)
    return bm25, tokenized

# ----- Ingestion helper -----

def ingest_document(
    session_id: str,
    filename: str,
    file_path: str,
    chunk_size: int = 500,
    overlap: int = 100,
    strategy: str = "fixed",
    progress_cb: Optional[Callable[[str], None]] = None,
) -> Tuple[int, int, Optional[str]]:
    """Process an uploaded document and store its vectors and BM25 index.
    Returns (page_count, chunk_count, title) where title is a best-effort
    document title (PDF /Title, DOCX core props, MD H1, TXT first line) or
    None when only the filename should be used for display.

    chunk_size/overlap/strategy are per-upload settings; the strategy that
    produced the index is recorded in each chunk's metadata so retrieval
    results stay explainable. progress_cb(stage) is called with real pipeline
    stages ("parsing", "chunking", "embedding", "indexing") so UIs can show
    honest progress instead of a timed animation.
    """
    if strategy not in CHUNKERS:
        raise ValueError(f"Unknown chunking strategy: {strategy}")

    ext = os.path.splitext(filename)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type '{ext}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")

    def report(stage: str):
        if progress_cb:
            progress_cb(stage)

    embedder_name = get_embedder().name
    all_chunks = []
    metadata = []

    def add_chunks(chunks: List[str], page_number: int, extra: Optional[dict] = None):
        for chunk in chunks:
            all_chunks.append(chunk)
            meta = {
                "filename": filename,
                "page_number": page_number,
                "chunk_strategy": strategy,
                "chunk_size": chunk_size,
                "overlap": overlap,
                "embedder": embedder_name,
            }
            if extra:
                meta.update(extra)
            metadata.append(meta)

    title = None
    page_count = 1
    if ext == ".pdf":
        report("parsing")
        pages = extract_text_from_pdf(file_path)
        title = extract_document_title(filename, file_path, pages=pages)
        chunker = CHUNKERS[strategy]
        for page_num, page_text in pages:
            add_chunks(chunker(page_text, chunk_size, overlap), page_num)
        page_count = len(pages)
    elif ext == ".txt":
        report("parsing")
        pages = extract_text_from_txt(file_path)
        title = extract_document_title(filename, file_path, pages=pages)
        chunker = CHUNKERS[strategy]
        for page_num, page_text in pages:
            add_chunks(chunker(page_text, chunk_size, overlap), page_num)
    elif ext == ".docx":
        report("parsing")
        blocks = extract_blocks_from_docx(file_path)
        title = extract_document_title(filename, file_path, blocks=blocks)
        add_chunks(chunk_blocks(blocks, chunk_size, overlap, structure_aware=(strategy == "structure_aware")), 1)
    elif ext == ".md":
        report("parsing")
        blocks = extract_blocks_from_markdown(file_path)
        title = extract_document_title(filename, file_path, blocks=blocks)
        add_chunks(chunk_blocks(blocks, chunk_size, overlap, structure_aware=(strategy == "structure_aware")), 1)
    elif ext == ".csv":
        report("parsing")
        csv_chunks, data_rows = extract_csv_chunks(file_path, chunk_size)
        report("chunking")
        for chunk_text, row_start, row_end in csv_chunks:
            add_chunks([chunk_text], 1, {"row_start": row_start, "row_end": row_end})
        page_count = data_rows  # page_count carries the data-row count for CSVs

    report("chunking")
    report("embedding")
    embeddings = embed_chunks(all_chunks)

    # Store in Chroma collection
    report("indexing")
    session = get_session_vectors(session_id)
    collection = session["collection"]
    ids = [f"{filename}_p{meta['page_number']}_c{i}" for i, meta in enumerate(metadata)]
    if title:
        for meta in metadata:
            meta["title"] = title
    collection.add(
        documents=all_chunks,
        embeddings=embeddings.tolist(),
        metadatas=metadata,
        ids=ids,
    )

    # Rebuild the session BM25 index over the FULL corpus so every document in the
    # session stays sparse-searchable (ingest is per-document, but the index is per-session).
    all_docs = collection.get(include=["documents"])
    bm25, tokenized = build_bm25_index(all_docs["documents"])
    persist_bm25(session_id, tokenized, all_docs["ids"])
    session["bm25"] = bm25
    session["bm25_ids"] = all_docs["ids"]

    return page_count, len(all_chunks), title
