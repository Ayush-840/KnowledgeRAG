"""Regenerate the binary documents in the eval corpus (PDF + DOCX).

The corpus ships with generated binaries so reviewers can regenerate them:
    .venv/bin/python -m eval.gen_binary_docs

The PDF intentionally includes running headers/footers and page numbers so
the Phase 0.2 extraction cleaner (detect_repeated_lines / clean_page_text)
has real layout artifacts to strip — the extraction baseline test guards
against regressions. The DOCX uses heading styles so the structure-aware
chunker splits along heading boundaries.
"""

from pathlib import Path

OUT = Path(__file__).resolve().parents[3] / "eval-docs"

PRODUCT_SHEET_SECTIONS = [
    ("What is Atlas?", (
        "Atlas is Aurora Labs' business intelligence platform for teams that "
        "query large datasets directly against their data warehouse. Atlas "
        "connects to Snowflake, BigQuery, Postgres, and Redshift, builds "
        "dashboards from live queries, and delivers scheduled reports by "
        "email or Slack. Atlas is designed for analysts who want the "
        "performance of a purpose-built query engine without managing "
        "infrastructure."
    )),
    ("Query engine", (
        "Atlas is powered by the Calcite query engine, a distributed SQL "
        "engine with a cost-based optimizer and vectorized in-memory "
        "execution. Queries are pushed down to the source system wherever "
        "possible, so a dashboard over a filtered subset of a large table "
        "does not pull the whole table. The engine maintains an in-memory "
        "columnar cache that makes warm queries run at p95 under 2 seconds "
        "on the reference cluster. Cold queries, which compile a fresh plan, "
        "measure at roughly 4 seconds p95; see the platform architecture "
        "report for details."
    )),
    ("Dashboards and reports", (
        "The dashboard builder supports 40+ chart types including bar, line, "
        "area, heatmap, funnel, and pivot tables. Dashboards are versioned "
        "and every edit is written to the audit log. Scheduled reports "
        "deliver a snapshot or a live link to email or Slack on a daily, "
        "weekly, or monthly cadence. Dashboards can be embedded in your own "
        "application with the embed SDK, which supports single sign-on and "
        "row-level security."
    )),
    ("Security and governance", (
        "Atlas encrypts data in transit with TLS 1.2+ and at rest with "
        "AES-256. Row-level security lets administrators constrain what each "
        "user or team can see without duplicating tables. SSO is supported "
        "via SAML and OIDC on Standard and Enterprise tiers. The audit log "
        "records every query run, dashboard edit, and permission change, and "
        "is retained for 12 months on Standard and 36 months on Enterprise."
    )),
    ("Pricing", (
        "Atlas is priced per user per month. Starter is $29, Standard is "
        "$49, and Enterprise is $99 per user per month with annual billing. "
        "Monthly billing is available at a 20% surcharge. Nonprofit and "
        "education discounts apply. See product-pricing.csv for the full "
        "price list, including bundles that combine Atlas with Beacon and "
        "Nimbus."
    )),
    ("Getting started", (
        "Create a workspace at app.auroralabs.example, connect a data "
        "source, and invite users. The free trial includes 14 days of "
        "Standard features with a 1 TB query allowance. The REST API is "
        "documented in api-reference.md, and the query editor includes a "
        "live explain plan and cost estimation. Support is available "
        "through the portal during business hours, and 24/7 support is "
        "included on Enterprise."
    )),
    ("Product SKUs", (
        "Each product and tier has a stable SKU used in orders and the "
        "billing system. The Atlas analytics platform is SKU ATL-100, "
        "with ATL-101 for the Enterprise tier and ATL-102 for Atlas "
        "embedded deployments. Beacon is SKU BCN-100, Nimbus is NMB-100, "
        "and the combined bundle is BND-300. SKUs appear on invoices and "
        "in the API's billing endpoints, and customers should reference "
        "the SKU when raising billing questions. SKUs do not change when "
        "prices change, so past invoices remain traceable."
    )),
    ("Data connectors", (
        "Atlas connects natively to Snowflake, BigQuery, Postgres, and "
        "Redshift, plus Amazon S3 for data lake queries. Connections are "
        "managed centrally: credentials are encrypted at rest and the "
        "connection pool handles warehouse session limits automatically. "
        "Connector settings include warehouse or dataset selection, "
        "connection timeouts, and optional read-only credentials. BigQuery "
        "connections can use service-account keys or workload identity, and "
        "Snowflake connections support key-pair authentication. The "
        "connector framework pushes predicates down to the source so "
        "filtered queries never scan the full table."
    )),
    ("Embedding and API access", (
        "Every Atlas capability is available through the REST API, which "
        "is documented in api-reference.md and covered by official SDKs "
        "for Python, TypeScript, and Go. Dashboards can be embedded in "
        "your application with the embed SDK, which supports row-level "
        "security, SSO, and custom theming. API keys are scoped per "
        "capability and rotated on a schedule you control. The audit log "
        "records every API call with the key identifier, so unusual usage "
        "is attributable to a specific credential."
    )),
    ("Admin console", (
        "Workspace administrators manage users, teams, data source "
        "connections, and billing from the admin console. Role-based "
        "access control assigns admin, editor, and viewer roles per team. "
        "The console supports SCIM provisioning for user lifecycle "
        "management, and session policies enforce login frequency and "
        "device requirements. The console also hosts the audit log, "
        "export tools, and the data subject request portal described in "
        "the privacy policy."
    )),
    ("Migration from legacy BI", (
        "Teams migrating from legacy business intelligence tools can use "
        "the migration assistant, which imports dashboard definitions and "
        "saved queries from Tableau, Looker, and Power BI. The assistant "
        "maps data source connections, converts chart types to the closest "
        "Atlas equivalent, and flags queries that need manual review. "
        "Migration runs are staged: review the converted assets in a "
        "staging workspace before promoting them, and the audit log keeps "
        "a record of the import. Most teams complete a migration in under "
        "two weeks."
    )),
    ("Performance expectations", (
        "Atlas performance depends on the connected warehouse and the "
        "query shape. On the reference cluster, warm dashboard queries run "
        "at p95 under 2 seconds, and the columnar cache serves repeated "
        "queries without contacting the warehouse. Scheduled report "
        "generation is capped at 200 reports per hour per workspace on "
        "Standard and 1,000 on Enterprise. The performance commitments "
        "and the availability guarantee are defined in sla-agreement.md, "
        "and measured uptime is published monthly on the status page."
    )),
]

SECURITY_POLICY_SECTIONS = [
    ("Purpose and scope", (
        "This Information Security Policy defines the security controls "
        "that protect the Aurora Labs platform and the customer data it "
        "processes. It applies to all employees, contractors, and systems "
        "that touch the production environment, including the Atlas, "
        "Beacon, and Nimbus services and the supporting infrastructure."
    )),
    ("Access control", (
        "Access to production systems is granted on a least-privilege "
        "basis and requires multi-factor authentication. Production access "
        "is requested through the access management tool and approved by "
        "the service owner. Access is reviewed quarterly, and revoked "
        "within 24 hours of an employee departing. Service accounts use "
        "short-lived credentials that rotate automatically every 90 days. "
        "Privileged access sessions are recorded and audited."
    )),
    ("Data protection", (
        "Customer data is encrypted at rest with AES-256 and in transit "
        "with TLS 1.2+. Backups are encrypted and stored in a separate "
        "region from the primary data. Data is retained per the privacy "
        "policy retention schedule and deleted within 30 days of the "
        "retention window expiring. Access to customer data by staff is "
        "limited to the support and engineering roles that require it, is "
        "logged, and is reviewed for anomalous patterns. API keys used by "
        "integrations must be rotated every 180 days, enforced by the key "
        "management service which revokes keys that exceed the rotation "
        "window."
    )),
    ("Change management", (
        "Changes to production follow a change management process. Code "
        "changes require a peer review and pass the automated test suite "
        "before merging. Deployments are automated and canary-based, with "
        "automatic rollback on error-rate thresholds. Infrastructure "
        "changes are made through infrastructure-as-code and reviewed in "
        "the same pull-request process as code. Emergency changes require "
        "a post-change review within 5 business days."
    )),
    ("Incident response", (
        "Security incidents are handled under the incident response "
        "process described in the incident runbook, with the security "
        "team leading. Suspected incidents are reported to the security "
        "on-call immediately. Notifications to customers and authorities "
        "follow the timelines in the GDPR FAQ and the DPA: customers "
        "within 72 hours of confirmation, and supervisory authorities "
        "where the GDPR requires it."
    )),
    ("Vendor management", (
        "Third-party vendors with access to customer data are reviewed "
        "before onboarding, including a security questionnaire and, for "
        "higher-risk vendors, review of their SOC 2 report. Vendors are "
        "re-reviewed annually, and the subprocessor list is published and "
        "updated at least 30 days before a new subprocessor is added."
    )),
    ("Training and awareness", (
        "All employees complete annual security and privacy training and "
        "acknowledge the acceptable use policy. Engineering teams run a "
        "quarterly game day that includes a security scenario. Phishing "
        "simulations run quarterly, and employees who fail twice in a "
        "year complete additional training."
    )),
    ("Monitoring and compliance", (
        "The security program is audited annually and the SOC 2 Type II "
        "report is available to customers under NDA. Monitoring covers "
        "authentication, API usage, and data access, with alerting through "
        "Beacon. The program is reviewed quarterly by the security "
        "committee, and this policy is reviewed annually."
    )),
    ("Workstation security", (
        "Company-managed workstations run a standard build with disk "
        "encryption enforced, automatic updates, and endpoint detection "
        "and response (EDR) agents. Software installation is restricted to "
        "an approved catalog. Employees must lock their screens when "
        "unattended and must not connect personal devices to the corporate "
        "network except through the approved VPN. Laptop loss or theft is "
        "reported to IT within 24 hours so the device can be remotely "
        "wiped and the account sessions revoked."
    )),
    ("Data classification and handling", (
        "Data is classified as public, internal, confidential, or "
        "restricted. Customer content in production systems is treated as "
        "confidential at minimum, and customer content that includes "
        "personal data is restricted. Confidential and restricted data are "
        "encrypted in transit and at rest, are only stored on approved "
        "systems, and are only shared through approved channels. Downloads "
        "of production data to personal devices are prohibited; analytics "
        "on production data run in the analytics environment with the "
        "same access controls as production."
    )),
    ("Acceptable use", (
        "Company systems are provided for business purposes. Occasional "
        "personal use is permitted if it does not interfere with work, "
        "consume excessive resources, or expose company data. Employees "
        "may not use company systems to store personal data of others "
        "without a business need, may not install unapproved software, and "
        "may not share accounts or credentials. Violations are handled "
        "through the disciplinary process and may result in termination "
        "for serious or repeated violations."
    )),
    ("Physical security", (
        "Aurora Labs offices use badge-controlled entry, visitor logging, "
        "and monitored server rooms. Production infrastructure runs in "
        "cloud provider facilities that meet the provider's SOC 2 and ISO "
        "27001 certifications; there is no on-premises production hardware. "
        "Employees working from home are expected to use the company VPN "
        "for any access to internal systems and to keep their workspace "
        "screen-locked when it is not in use. Tailgating and unescorted "
        "visitors in secure areas are reported to the facilities team. "
        "Office security incidents are reported to People Operations, who "
        "coordinate with the security team."
    )),
    ("Backup and recovery", (
        "Customer data is backed up continuously with hourly point-in-time "
        "recovery for 30 days and daily backups retained for 12 months. "
        "Backups are encrypted and stored in a separate region from the "
        "primary data, and restore is tested quarterly through a scheduled "
        "restore drill that validates row counts and checksums. Recovery "
        "time objectives are documented per service: Atlas recovers in "
        "under 4 hours, Nimbus in under 2 hours, and Marina in under 1 "
        "hour from the point-in-time store. The backup policy is reviewed "
        "annually and when the platform architecture changes materially."
    )),
    ("Security roles and responsibilities", (
        "The Chief Information Security Officer owns this policy and the "
        "security program. The security team operates monitoring, incident "
        "response, and vendor review. Engineering managers are responsible "
        "for the security of the systems their teams build, including code "
        "review, dependency updates, and threat modeling of new features. "
        "Every employee is responsible for following this policy and "
        "reporting suspected incidents. The security committee meets "
        "monthly to review metrics, open corrective actions, and changes "
        "to the threat landscape, and approves any exception to this "
        "policy before it is granted. Exceptions are time-boxed and "
        "tracked in the security register."
    )),
]


def build_pdf(path: Path) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    def header_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColorRGB(0.4, 0.4, 0.4)
        canvas.drawCentredString(letter[0] / 2, letter[1] - 0.5 * inch, "Aurora Labs Product Sheet — Atlas")
        canvas.drawCentredString(letter[0] / 2, 0.5 * inch, f"— {doc.page} —")
        canvas.restoreState()

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], spaceAfter=14)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=10.5, leading=15, spaceAfter=10)

    doc = SimpleDocTemplate(str(path), pagesize=letter, title="Atlas Product Sheet",
                            author="Aurora Labs")
    story = [Paragraph("Atlas Product Sheet", h1)]
    for title, text in PRODUCT_SHEET_SECTIONS:
        story.append(Paragraph(title, styles["Heading2"]))
        for para in text.split("\n\n"):
            story.append(Paragraph(para, body))
        story.append(Spacer(1, 8))
        story.append(PageBreak())
    doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer)


def build_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.core_properties.title = "Aurora Labs Information Security Policy"
    doc.core_properties.author = "Aurora Labs Security Team"
    doc.add_heading("Aurora Labs Information Security Policy", level=0)
    doc.add_paragraph(
        "Effective 2025-03-01. Owner: Chief Information Security Officer. "
        "This policy is reviewed annually and on any material change to the platform."
    )
    for title, text in SECURITY_POLICY_SECTIONS:
        doc.add_heading(title, level=1)
        for para in text.split("\n\n"):
            doc.add_paragraph(para)
    doc.save(path)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    build_pdf(OUT / "atlas-product-sheet.pdf")
    build_docx(OUT / "security-policy.docx")
    print(f"regenerated {OUT / 'atlas-product-sheet.pdf'} and {OUT / 'security-policy.docx'}")


if __name__ == "__main__":
    main()
