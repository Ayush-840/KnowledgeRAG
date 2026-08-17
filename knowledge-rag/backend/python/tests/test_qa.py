"""Automated QA suite — maps 1:1 to QA_TEST_PLAN.md scenarios S1-S13.

Run from backend/python:  .venv/bin/python -m pytest tests/ -v
Hermetic: stub embedder (no model loads), reranker disabled, temp persistence
and log dirs. Never touches the network or your .env credentials.
"""

import json
import os
from pathlib import Path

import pytest

from app.dependencies import get_session_vectors
from app.llm import verify_citations
from app.retrieval import reciprocal_rank_fusion
from eval import llm as eval_llm
from app.utils import (
    clean_extracted_pages,
    clean_page_text,
    detect_repeated_lines,
    extract_csv_chunks,
    extract_blocks_from_markdown,
    extract_text_from_pdf,
)

SAMPLE_DIR = Path(__file__).resolve().parents[3] / "sample-docs"

TXT_SAMPLE = "Aurora Labs is a software company founded in 2019. " \
             "Atlas costs $49 per user per month. Beacon integrates with Slack."
MD_SAMPLE = "# Aurora Labs FAQ\n\n## Billing\n\nAtlas costs $49 per user per month.\n\n## Security\n\nData is encrypted with AES-256."
CSV_SAMPLE = "Product,Category,Price\nAtlas,Analytics,49\nBeacon,Alerting,19\nNimbus,Pipeline,0\n"
DOCX_HEADINGS = ["Getting started", "Billing", "Security", "Integrations"]


# ---------- helpers ----------

def _ingest(client, session_id: str, filename: str, content: bytes, **params):
    mime = "application/octet-stream"
    return client.post(
        f"/ingest/{session_id}",
        files={"file": (filename, content, mime)},
        params=params,
    )


@pytest.fixture()
def session_id():
    import uuid

    return f"qa-{uuid.uuid4().hex[:10]}"


@pytest.fixture()
def ready_session(client, session_id):
    """A session with the TXT sample ingested (BM25 initialized)."""
    r = _ingest(client, session_id, "sample.txt", TXT_SAMPLE.encode())
    assert r.status_code == 200, r.text
    return session_id


# ---------- S0: PDF extraction hygiene (Phase 0.2) ----------

def test_pdf_extraction_strips_page_numbers():
    text = "Page 3\n\nAurora Labs was founded in 2019.\n\n4\n"
    assert clean_page_text(text) == "Aurora Labs was founded in 2019."


def test_pdf_extraction_splits_glued_numeric_prefix():
    # The "04Build" artifact: a short numeric page marker glued onto the first
    # word of the page with no whitespace. Digits dropped, word kept.
    text = "04Build pipelines in production.\n\nThe Atlas platform costs $49."
    cleaned = clean_page_text(text)
    assert cleaned.startswith("Build pipelines in production.")
    # Legitimate numbers inside body text are untouched
    assert "$49" in cleaned


def test_pdf_extraction_keeps_legitimate_leading_numbers():
    # A 4-digit year at line start must NOT be treated as a glued page marker
    text = "2024 revenue grew 40% year over year."
    assert clean_page_text(text) == "2024 revenue grew 40% year over year."


def test_pdf_extraction_glue_strip_survives_header_line():
    # A running header above the content must not consume the glue-strip:
    # "04Build" on the first content line is still split, even though the
    # header line comes first.
    pages = [
        (1, "Aurora Labs Confidential\n\nAtlas deployment guide.\n\nPage 1\nAurora Labs Confidential"),
        (2, "Aurora Labs Confidential\n\n04Build pipelines for prod.\n\nPage 2\nAurora Labs Confidential"),
        (3, "Aurora Labs Confidential\n\n05Runbook for incidents.\n\nPage 3\nAurora Labs Confidential"),
        (4, "Aurora Labs Confidential\n\nMore content.\n\nPage 4\nAurora Labs Confidential"),
    ]
    cleaned = clean_extracted_pages(pages)
    assert cleaned[1][1].startswith("Build pipelines for prod.")
    assert cleaned[2][1].startswith("Runbook for incidents.")
    assert "Aurora Labs Confidential" not in cleaned[1][1]


def test_pdf_extraction_strips_repeated_headers_footers():
    pages = [
        (1, "Aurora Labs\n\nIntro page content.\n\nPage 1\nAurora Labs"),
        (2, "Aurora Labs\n\nMore content here.\n\nPage 2\nAurora Labs"),
        (3, "Aurora Labs\n\nFinal content.\n\nPage 3\nAurora Labs"),
        (4, "Aurora Labs\n\nMore content.\n\nPage 4\nAurora Labs"),
    ]
    repeated = detect_repeated_lines(pages)
    assert "Aurora Labs" in repeated  # running header/footer detected
    cleaned = clean_extracted_pages(pages)
    # The boilerplate header/footer is gone from every page; body text survives
    for _, text in cleaned:
        assert "Aurora Labs\n\n" not in text
        assert "content" in text
    assert "Intro page content." in cleaned[0][1]


def test_pdf_extraction_baseline(client):
    """Regression check: re-extract the sample corpus and diff against the stored
    baseline so a future parser change that reintroduces layout contamination is
    caught here, not by a user. Regenerate with UPDATE_BASELINES=1."""
    baseline_path = Path(__file__).parent / "baselines" / "pdf_extraction.json"
    if not baseline_path.exists():
        pytest.skip("extraction baseline not present")
    baseline = json.loads(baseline_path.read_text(encoding="utf-8"))
    for filename, record in baseline.items():
        pdf = SAMPLE_DIR / filename
        if not pdf.exists():
            continue
        pages = extract_text_from_pdf(str(pdf))
        actual = [{"page": p, "text": t} for p, t in pages]
        if os.getenv("UPDATE_BASELINES") == "1":
            baseline[filename] = {"pages": actual}
            continue
        assert actual == record["pages"], (
            f"PDF extraction for {filename} changed! If the change is intended, "
            "regenerate with UPDATE_BASELINES=1."
        )
    if os.getenv("UPDATE_BASELINES") == "1":
        baseline_path.write_text(json.dumps(baseline, indent=2, ensure_ascii=False), encoding="utf-8")


# ---------- S1: extension whitelist ----------

def test_s1_unsupported_extension_rejected(client, session_id):
    r = _ingest(client, session_id, "malware.exe", b"MZ...")
    assert r.status_code == 400
    assert "Unsupported file type" in r.json()["detail"]
    assert ".exe" in r.json()["detail"]


def test_s1_png_rejected(client, session_id):
    r = _ingest(client, session_id, "image.png", b"\x89PNG")
    assert r.status_code == 400


# ---------- S2: size guardrail ----------

def test_s2_oversize_rejected(client, session_id, monkeypatch):
    import app.routes as routes

    monkeypatch.setattr(routes, "MAX_UPLOAD_BYTES", 100)
    r = _ingest(client, session_id, "big.txt", b"x" * 5000)
    assert r.status_code == 413
    assert "upload limit" in r.json()["detail"]


def test_s2_within_limit_ok(client, session_id):
    r = _ingest(client, session_id, "small.txt", b"hello world")
    assert r.status_code == 200


# ---------- S3: CSV row limit ----------

def test_s3_csv_row_limit(client, session_id, monkeypatch):
    import app.utils as utils

    monkeypatch.setattr(utils, "MAX_CSV_ROWS", 3)
    rows = "\n".join([f"r{i}" for i in range(10)])
    r = _ingest(client, session_id, "too_many.csv", f"h\n{rows}".encode())
    assert r.status_code == 400
    assert "row limit" in r.json()["detail"]


# ---------- S4: multi-format ingestion ----------

@pytest.mark.parametrize(
    "filename,content",
    [
        ("doc.txt", TXT_SAMPLE.encode()),
        ("doc.md", MD_SAMPLE.encode()),
        ("doc.csv", CSV_SAMPLE.encode()),
    ],
)
def test_s4_text_formats_ingest(client, session_id, filename, content):
    r = _ingest(client, session_id, filename, content)
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["chunk_count"] >= 1
    assert body["chunk_strategy"] == "fixed"


def test_s4_docx_ingests(client, session_id, tmp_path):
    from docx import Document

    doc_path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_heading("Guide", level=1)
    doc.add_paragraph("This is the body of the guide document.")
    doc.save(doc_path)
    r = _ingest(client, session_id, "doc.docx", doc_path.read_bytes())
    assert r.status_code == 200, r.text
    assert r.json()["chunk_count"] >= 1


def test_s4_pdf_ingests(client, session_id):
    pdf = SAMPLE_DIR / "aurora-labs.pdf"
    if not pdf.exists():
        pytest.skip("sample PDF not present")
    r = _ingest(client, session_id, "aurora-labs.pdf", pdf.read_bytes())
    assert r.status_code == 200, r.text
    assert r.json()["page_count"] >= 1


# ---------- S5: chunking strategy + settings validation ----------

def test_s5_structure_aware_uses_headings(client, session_id):
    r = _ingest(
        client, session_id, "faq.md", MD_SAMPLE.encode(), strategy="structure_aware"
    )
    assert r.status_code == 200, r.text
    # Each chunk should carry its heading line (verified via the document endpoint)
    dr = client.get(f"/documents/{session_id}/faq.md")
    assert dr.status_code == 200
    chunk_texts = [c["text"] for c in dr.json()["chunks"]]
    joined = "\n".join(chunk_texts)
    assert "Billing" in joined and "Security" in joined
    assert "Atlas costs $49" in joined


def test_s5_invalid_chunk_size_rejected(client, session_id):
    r = _ingest(client, session_id, "doc.txt", TXT_SAMPLE.encode(), chunk_size=10)
    assert r.status_code == 400


def test_s5_invalid_overlap_rejected(client, session_id):
    r = _ingest(client, session_id, "doc.txt", TXT_SAMPLE.encode(), overlap=500, chunk_size=200)
    assert r.status_code == 400


def test_s5_unknown_strategy_rejected(client, session_id):
    r = _ingest(client, session_id, "doc.txt", TXT_SAMPLE.encode(), strategy="magic")
    assert r.status_code == 400


def test_s5_markdown_block_extraction(tmp_path):
    md_file = tmp_path / "sample.md"
    md_file.write_text(MD_SAMPLE, encoding="utf-8")
    blocks = extract_blocks_from_markdown(str(md_file))
    headings = [b["heading"] for b in blocks if b.get("heading")]
    assert "Billing" in headings and "Security" in headings


# ---------- S6: CSV coherence ----------

def test_s6_csv_stays_tabular(client, session_id):
    r = _ingest(client, session_id, "products.csv", CSV_SAMPLE.encode())
    assert r.status_code == 200, r.text
    dr = client.get(f"/documents/{session_id}/products.csv")
    assert dr.status_code == 200
    chunks = dr.json()["chunks"]
    assert len(chunks) >= 1
    first = chunks[0]
    # Header row preserved and row range recorded
    assert "Product" in first["text"] and "Price" in first["text"]
    assert first["row_start"] == 2
    assert first["row_end"] >= 2


def test_s6_csv_chunker_units(tmp_path):
    csv_file = tmp_path / "p.csv"
    csv_file.write_text(CSV_SAMPLE, encoding="utf-8")
    chunks, data_rows = extract_csv_chunks(str(csv_file), max_words=500)
    assert data_rows == 3
    assert len(chunks) == 1
    text, row_start, row_end = chunks[0]
    assert "Product" in text and "Atlas" in text
    assert (row_start, row_end) == (2, 4)


# ---------- S7: hybrid scores labeled ----------

def test_s7_search_returns_labeled_scores(client, ready_session):
    r = client.post(f"/search/{ready_session}", json={"query": "Atlas price"})
    assert r.status_code == 200, r.text
    results = r.json()["results"]
    assert results, "expected at least one retrieved chunk"
    for chunk in results:
        scores = chunk["retrieval_scores"]
        # All four stage keys present; never blended into one number
        assert set(scores.keys()) == {"dense_similarity", "bm25", "rrf", "rerank"}
        assert scores["rrf"] is not None
        assert scores["rerank"] is None  # reranker disabled in tests


# ---------- S8: RRF fusion math ----------

def test_s8_rrf_symmetric_ranks():
    lists = [["a", "b", "c"], ["b", "a", "d"]]
    fused = reciprocal_rank_fusion(lists, k=60)
    # a: rank 1 + rank 2 -> 1/61 + 1/62 ; b: rank 2 + rank 1 -> same
    assert abs(fused["a"] - fused["b"]) < 1e-12
    assert abs(fused["a"] - (1 / 61 + 1 / 62)) < 1e-12


def test_s8_rrf_k_constant():
    lists = [["a", "b"], ["b", "a"]]
    assert reciprocal_rank_fusion(lists, k=60)["a"] > reciprocal_rank_fusion(lists, k=1000)["a"]


# ---------- S9: retrieval funnel ----------

def test_s9_funnel_counts(client, ready_session):
    r = client.post(f"/search/{ready_session}", json={"query": "Atlas price"})
    assert r.status_code == 200
    body = r.json()
    assert body["candidates_retrieved"] >= body["candidates_sent_to_llm"]
    assert 1 <= body["candidates_sent_to_llm"] <= 5  # RERANK_TOP_K default


# ---------- S10: context isolation ----------

def test_s10_sessions_isolated(client, session_id):
    # Two sessions, different documents
    r_a = _ingest(client, session_id, "faq.md", MD_SAMPLE.encode())
    assert r_a.status_code == 200
    r_b = _ingest(client, f"{session_id}-b", "products.csv", CSV_SAMPLE.encode())
    assert r_b.status_code == 200

    # Searching session A must only surface faq.md chunks
    ra = client.post(f"/search/{session_id}", json={"query": "Atlas"})
    assert ra.status_code == 200
    assert all(c["filename"] == "faq.md" for c in ra.json()["results"])

    # And session B only products.csv chunks
    rb = client.post(f"/search/{session_id}-b", json={"query": "Beacon"})
    assert rb.status_code == 200
    assert all(c["filename"] == "products.csv" for c in rb.json()["results"])


# ---------- S11: citation verification ----------

def test_s11_strips_fabricated_markers():
    context = [
        {"id": "c1", "filename": "a.md", "page_number": 1, "text": "Atlas costs $49."},
        {"id": "c2", "filename": "a.md", "page_number": 1, "text": "Beacon integrates with Slack."},
    ]
    answer, citations = verify_citations("Atlas is $49 [1]. Beacon uses Slack [2]. Also [9] and [2] again.", context)
    assert "[1]" in answer and "[2]" in answer
    assert "[9]" not in answer
    assert len(citations) == 2
    assert [c["marker"] for c in citations] == [1, 2]
    assert citations[0]["id"] == "c1" and citations[1]["id"] == "c2"


def test_s11_no_citations_when_none_valid():
    answer, citations = verify_citations("No sources here [3] [42].", [{"id": "c1", "text": "x"}])
    assert answer == "No sources here  ."
    assert citations == []


# ---------- S11b: eval response-structure pre-check ----------

COMPLIANT_ANSWER = (
    "**Executive Summary / Core Answer**\n"
    "Two plus two is four [1].\n\n"
    "**Key Takeaways & Concepts**\n"
    "- Addition is commutative [1].\n\n"
    "**Citation References**\n"
    "- [1] Math Notes, page 1"
)


def test_s11b_structure_check_accepts_compliant_answer():
    ok, issues = eval_llm.check_structure(COMPLIANT_ANSWER, 3)
    assert ok, issues
    assert issues == []


def test_s11b_structure_check_empty_answer():
    ok, issues = eval_llm.check_structure("", 3)
    assert not ok
    assert "empty" in issues[0]


def test_s11b_structure_check_missing_sections():
    ok, issues = eval_llm.check_structure("**Key Takeaways & Concepts**\n- stuff [1]", 3)
    assert not ok
    joined = " | ".join(issues)
    assert "Executive Summary" in joined
    assert "Citation References" in joined


def test_s11b_structure_check_unresolved_body_marker():
    answer = (
        "**Executive Summary**\nFact [2].\n\n"
        "**Key Takeaways**\n- point\n\n"
        "**Citation References**\n- [1] Doc, page 1"
    )
    ok, issues = eval_llm.check_structure(answer, 3)
    assert not ok
    assert any("marker [2]" in i and "missing" in i for i in issues)


def test_s11b_structure_check_out_of_range_marker():
    answer = (
        "**Executive Summary**\nFact [9].\n\n"
        "**Key Takeaways**\n- point\n\n"
        "**Citation References**\n- [9] Doc, page 1"
    )
    ok, issues = eval_llm.check_structure(answer, 3)
    assert not ok
    assert any("exceeds context size" in i for i in issues)


# ---------- S12: graceful degradation ----------

def test_s12_chat_graceful_fallback_without_key(client, ready_session, monkeypatch):
    monkeypatch.delenv("OPENROUTER_API_KEY", raising=False)
    monkeypatch.delenv("NVIDIA_API_KEY", raising=False)
    r = client.post(f"/chat/{ready_session}", json={"query": "Atlas price"})
    assert r.status_code == 200
    data = r.json()
    assert "answer" in data and len(data["answer"]) > 0
    assert "citations" in data


def test_s12_chat_graceful_fallback_on_generation_failure(client, ready_session, monkeypatch):
    import app.routes as routes

    monkeypatch.setattr(routes.llm_client, "llm_available", lambda: True)
    monkeypatch.setattr(
        routes.llm_client, "generate_answer", lambda q, c: (_ for _ in ()).throw(RuntimeError("boom"))
    )
    r = client.post(f"/chat/{ready_session}", json={"query": "Atlas price"})
    assert r.status_code == 200
    data = r.json()
    assert "answer" in data and len(data["answer"]) > 0
    assert "citations" in data


# ---------- S13: observability / JSONL logging ----------

def test_s13_search_writes_jsonl(client, ready_session):
    r = client.post(f"/search/{ready_session}", json={"query": "Atlas price"})
    assert r.status_code == 200

    log_dir = os.getenv("QUERY_LOG_DIR")
    log_path = Path(log_dir) / "queries.jsonl"
    assert log_path.exists(), "query log file not written"

    records = [json.loads(line) for line in log_path.read_text().splitlines() if line.strip()]
    last = records[-1]
    assert last["session_id"] == ready_session
    assert last["query"] == "Atlas price"
    assert "latency_ms" in last
    assert {"dense", "bm25", "fusion", "rerank", "total"} <= set(last["latency_ms"].keys())
    assert last["retrieved"]  # fused pool with stage scores
    assert last["final_answer"] is None  # search has no generation stage


# ---------- S14: real ingestion stages + document titles + chat titles ----------

def test_s14_ingest_streams_real_stages(client, session_id):
    """?stream=1 returns SSE events for real pipeline stages (parsing -> chunking
    -> embedding -> indexing -> done), not a timed animation."""
    r = _ingest(client, session_id, "stages.md", MD_SAMPLE.encode(), stream=True)
    assert r.status_code == 200
    assert r.headers.get("content-type", "").startswith("text/event-stream")
    events = []
    for line in r.text.splitlines():
        if line.startswith("data: "):
            events.append(json.loads(line[6:]))
    stages = [e["stage"] for e in events]
    assert stages == ["parsing", "chunking", "embedding", "indexing", "done"]
    done = events[-1]
    assert done["result"]["chunk_count"] >= 1
    assert done["result"]["title"] == "Aurora Labs FAQ"


def test_s14_ingest_stream_reports_errors(client, session_id):
    """Mid-stream failures surface as an error event, not a broken stream."""
    # An empty CSV passes validation but fails inside ingest_document
    r = _ingest(client, session_id, "empty.csv", b"\n", stream=True)
    assert r.status_code == 200
    events = [
        json.loads(line[6:])
        for line in r.text.splitlines()
        if line.startswith("data: ")
    ]
    assert events[-1]["stage"] == "error"
    assert "empty" in events[-1]["error"].lower()


def test_s14_title_endpoint_heuristic(client):
    """/title returns a deterministic heuristic title when no LLM key is set."""
    r = client.post("/title", json={"query": "What are the key findings in the Q3 budget report?"})
    assert r.status_code == 200
    title = r.json()["title"]
    assert title and len(title) <= 48
    assert title == "What are the key findings in the Q3"


def test_s14_title_endpoint_requires_query(client):
    r = client.post("/title", json={"query": "   "})
    assert r.status_code == 400


def test_s14_document_title_in_metadata(client, session_id):
    """Ingest prefers a document's own title (MD H1 here) and exposes it in the
    response, the /documents list, and chunk metadata."""
    r = _ingest(client, session_id, "faq.md", MD_SAMPLE.encode())
    assert r.status_code == 200, r.text
    assert r.json()["title"] == "Aurora Labs FAQ"

    listing = client.get(f"/documents/{session_id}").json()["documents"]
    assert listing[0]["title"] == "Aurora Labs FAQ"

    # Chunks carry the title so search results/citations can display it
    doc = client.get(f"/documents/{session_id}/faq.md").json()
    assert doc["title"] == "Aurora Labs FAQ"

    session = get_session_vectors(session_id)
    meta = session["collection"].get(include=["metadatas"])["metadatas"][0]
    assert meta.get("title") == "Aurora Labs FAQ"


def test_s14_plain_txt_first_line_title(client, session_id):
    """A .txt whose first line is short becomes the title."""
    content = "Onboarding Guide\n\nWeek one: set up your laptop and accounts.\n"
    r = _ingest(client, session_id, "notes.txt", content.encode())
    assert r.status_code == 200, r.text
    assert r.json()["title"] == "Onboarding Guide"


def test_s14_csv_has_no_title(client, session_id):
    """CSVs have no document title — callers fall back to the filename."""
    r = _ingest(client, session_id, "products.csv", CSV_SAMPLE.encode())
    assert r.status_code == 200, r.text
    assert r.json()["title"] is None


def test_s14_chroma_cosine_space(client, session_id):
    """Collections are created with hnsw:space=cosine so the dense_similarity
    label (1 - distance) is an actual cosine similarity, not L2."""
    _ingest(client, session_id, "sample.txt", TXT_SAMPLE.encode())
    collection = get_session_vectors(session_id)["collection"]
    assert collection.metadata.get("hnsw:space") == "cosine"


def test_s14_search_carries_title(client, session_id):
    r = _ingest(client, session_id, "faq.md", MD_SAMPLE.encode())
    assert r.status_code == 200
    sr = client.post(f"/search/{session_id}", json={"query": "Atlas"})
    assert sr.status_code == 200
    for chunk in sr.json()["results"]:
        assert chunk["title"] == "Aurora Labs FAQ"


# ---------- S15: 3D vector space explorer ----------

def test_s15_space_projection(client, ready_session):
    """GET /space returns compact 3D points for every chunk in the session."""
    r = client.get(f"/space/{ready_session}")
    assert r.status_code == 200
    body = r.json()
    assert body["point_count"] >= 1
    assert body["method"] in ("umap", "pca")
    assert body["clustered"] is False
    assert len(body["points"]) == body["point_count"]
    for p in body["points"]:
        assert {"id", "x", "y", "z", "filename"} <= set(p.keys())
        assert isinstance(p["x"], (int, float))


def test_s15_space_query_drop_in(client, ready_session):
    """A query transforms into the existing map (no re-layout) and reports the
    retrieval funnel's promoted vs. retrieved ids."""
    r = client.post(f"/space/{ready_session}/query", json={"query": "Atlas price"})
    assert r.status_code == 200
    body = r.json()
    assert {"x", "y", "z"} <= set(body["point"].keys())
    assert body["promoted_ids"], "expected promoted (top-k) chunk ids"
    assert len(body["promoted_ids"]) <= 5  # RERANK_TOP_K default
    assert body["retrieved_ids"]
    assert len(body["retrieved_ids"]) >= len(body["promoted_ids"])
    assert set(body["promoted_ids"]) <= set(body["retrieved_ids"])


def test_s15_space_query_requires_query(client, ready_session):
    r = client.post(f"/space/{ready_session}/query", json={"query": "  "})
    assert r.status_code == 400


def test_s15_space_empty_session_404(client, session_id):
    r = client.get(f"/space/{session_id}")
    assert r.status_code == 404


def test_s15_space_clustering_threshold(client, ready_session, monkeypatch):
    """Above UMAP_CLUSTER_THRESHOLD the projection degrades to representative
    cluster markers instead of a laggy full point cloud."""
    import app.space as space_mod

    monkeypatch.setattr(space_mod, "UMAP_CLUSTER_THRESHOLD", 0)
    r = client.get(f"/space/{ready_session}?force=1")
    assert r.status_code == 200
    body = r.json()
    assert body["clustered"] is True
    assert len(body["points"]) <= body["point_count"]
    for p in body["points"]:
        assert "count" in p and p["count"] >= 1


def test_s15_space_projection_cached(client, ready_session):
    """The fit is cached per document set: two calls return identical
    coordinates (no jitter), and only an added document recomputes."""
    a = client.get(f"/space/{ready_session}").json()
    b = client.get(f"/space/{ready_session}").json()
    assert a["points"] == b["points"]


def test_s16_sample_docs_uses_archived_golden(monkeypatch, tmp_path):
    """run_eval auto-switches to the archived 16-query golden set when docs_dir
    is sample-docs, so the README regression baseline stays reproducible
    (a 72-query eval-docs golden set against the tiny sample corpus would
    score near-zero recall)."""
    import eval.run_eval as run_eval

    docs = tmp_path / "sample-docs"
    docs.mkdir()
    (docs / "a.txt").write_text("Aurora Labs analytics platform overview.")

    default_golden = Path(run_eval.__file__).parent / "golden_set.json"
    sample_golden = Path(run_eval.__file__).parent / "golden_set_sample.json"
    assert sample_golden.exists(), "archived sample golden set missing"

    monkeypatch.setattr(run_eval, "DEFAULT_GOLDEN", default_golden)

    # sample-docs + default golden -> archived set
    resolved = run_eval.resolve_golden(docs, default_golden)
    assert resolved == sample_golden

    # explicit --golden is never overridden
    explicit = docs / "my-golden.json"
    explicit.write_text("[]")
    assert run_eval.resolve_golden(docs, explicit) == explicit

    # non-sample docs dir keeps the default
    other = tmp_path / "other"
    other.mkdir()
    assert run_eval.resolve_golden(other, default_golden) == default_golden


def test_s16_eval_session_reset(client, session_id):
    """Re-running the same eval settings combo must start from a clean
    collection: the persistent Chroma dir otherwise accumulates duplicate
    chunks, silently skewing recall denominators and breaking the README's
    reproducible baselines (regression guard for the reset run_eval now calls)."""
    from app import dependencies as deps

    _ingest(client, session_id, "sample.txt", TXT_SAMPLE.encode())
    sess = deps.get_session_vectors(session_id)
    assert len(sess["collection"].get()["ids"]) == 1
    assert sess["bm25"] is not None

    deps.reset_session(session_id)
    assert session_id not in deps.SESSION_REGISTRY

    fresh = deps.get_session_vectors(session_id)
    assert fresh["collection"].get()["ids"] == []
    assert fresh["bm25"] is None
